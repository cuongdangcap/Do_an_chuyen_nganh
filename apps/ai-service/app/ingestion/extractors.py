from __future__ import annotations

import re
import os
from pathlib import Path

from app.models.schemas import ExtractedSegment


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_segments(file_path: Path, file_name: str) -> list[ExtractedSegment]:
    extension = file_path.suffix.lower().lstrip(".")
    if extension == "pdf":
        return extract_pdf(file_path)
    if extension == "docx":
        return extract_docx(file_path)
    if extension in {"png", "jpg", "jpeg"}:
        return extract_image(file_path)
    if extension in {"txt", "md"}:
        return extract_text_file(file_path)
    raise ValueError(f"Unsupported file type: {extension or file_name}")


def extract_pdf(file_path: Path) -> list[ExtractedSegment]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if text:
            segments.append(ExtractedSegment(text=text, page_number=index))
    if not segments:
        raise ValueError("PDF does not contain extractable text. OCR is required for scanned PDFs.")
    return segments


def extract_docx(file_path: Path) -> list[ExtractedSegment]:
    from docx import Document

    document = Document(str(file_path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                blocks.append(" | ".join(cells))

    text = clean_text("\n\n".join(blocks))
    if not text:
        raise ValueError("DOCX does not contain extractable text.")
    return [ExtractedSegment(text=text)]


def extract_image(file_path: Path) -> list[ExtractedSegment]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise ValueError("OCR dependencies are not installed.") from exc

    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    elif Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe").exists():
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    tessdata_dir = os.environ.get("TESSDATA_DIR")
    if tessdata_dir:
        os.environ.setdefault("TESSDATA_PREFIX", tessdata_dir)
    config = f"--tessdata-dir {tessdata_dir}" if tessdata_dir else ""

    try:
        text = clean_text(pytesseract.image_to_string(Image.open(file_path), lang="vie+eng", config=config))
    except Exception as exc:  # noqa: BLE001
        raise ValueError("OCR failed. Check that Tesseract OCR and Vietnamese language data are installed.") from exc

    if not text:
        raise ValueError("OCR produced no text.")
    return [ExtractedSegment(text=text, page_number=1)]


def extract_text_file(file_path: Path) -> list[ExtractedSegment]:
    text = clean_text(file_path.read_text(encoding="utf-8", errors="ignore"))
    if not text:
        raise ValueError("Text file is empty.")
    return [ExtractedSegment(text=text)]
